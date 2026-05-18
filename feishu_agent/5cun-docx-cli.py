#!/usr/bin/env python3
"""
5寸照片 Word 文档生成工具
将图片转换为 5 寸排版的 Word 文档
"""
import sys
import os
from pathlib import Path

def create_5cun_docx(image_dir, output_path):
    """生成5寸照片Word文档"""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 设置页面为 5 寸 x 3.5 寸（横向）
    section = doc.sections[0]
    section.page_width = Inches(5)
    section.page_height = Inches(3.5)
    section.left_margin = Inches(0.25)
    section.right_margin = Inches(0.25)
    section.top_margin = Inches(0.25)
    section.bottom_margin = Inches(0.25)

    # 支持的图片格式
    image_exts = ['.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp']
    images = []

    image_dir_path = Path(image_dir)
    for ext in image_exts:
        images.extend(sorted(image_dir_path.glob(f'*{ext}')))
        images.extend(sorted(image_dir_path.glob(f'*{ext.upper()}')))

    if not images:
        print("错误：未找到图片文件")
        return False

    # 每行放2张照片
    for i in range(0, len(images), 2):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 左边的照片
        if i < len(images):
            run = paragraph.add_run()
            run.add_picture(str(images[i]), width=Inches(2.2))

        # 右边的照片
        if i + 1 < len(images):
            run = paragraph.add_run()
            run.add_picture(str(images[i + 1]), width=Inches(2.2))

    # 如果只有一张图，添加空行占位
    if len(images) == 1:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(str(images[0]), width=Inches(2.2))

    doc.save(output_path)
    print(f"SUCCESS: {output_path}")
    return True

def main():
    if len(sys.argv) < 3:
        print("用法: python 5cun-docx-cli.py <图片目录> <输出文档路径>")
        sys.exit(1)

    image_dir = sys.argv[1]
    output_path = sys.argv[2]

    if not os.path.isdir(image_dir):
        print(f"错误：目录不存在: {image_dir}")
        sys.exit(1)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    success = create_5cun_docx(image_dir, output_path)
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()
